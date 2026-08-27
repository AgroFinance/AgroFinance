"""Casos borde del parser Python que no dependen de paridad con TS.

Regresión de un bug real: un CSV exportado por Excel en configuración
regional ES/PE usa ';' como separador (la coma es el separador decimal
ahí). Antes del fix, csv.reader() con el dialecto por defecto leía cada
fila como una sola columna y el archivo caía en "no se encontró ninguna
columna numérica" pese a tener datos válidos.
"""

from pathlib import Path

from engine.parse_archivo import parsear_archivo

FIXTURES = Path(__file__).parent / "fixtures"


def test_csv_con_delimitador_punto_y_coma():
    ruta = FIXTURES / "bitacora_riego_punto_y_coma.csv"
    parseado = parsear_archivo(str(ruta), "csv")

    assert parseado.columnas == [
        "FUNDO", "PARCELA", "FECHA", "HORAS_BOMBEO", "TIPO_RIEGO",
        "VOLUMEN_M3", "OPERARIO", "NOTA",
    ]
    # 3 filas x 2 columnas numericas (HORAS_BOMBEO, VOLUMEN_M3)
    assert len(parseado.lineas) == 6
    campos = {l.campo_leido for l in parseado.lineas}
    assert campos == {"HORAS_BOMBEO", "VOLUMEN_M3"}

    primera = next(l for l in parseado.lineas if l.campo_leido == "VOLUMEN_M3")
    assert primera.valor == 57.4
    # Sin esto, VOLUMEN_M3 se lee como numero pero con unidad vacia y nunca
    # llega a la huella hidrica (que exige m3 o litros para reconocer agua).
    assert primera.unidad == "m3"


# ============================================================
# Números con separador de miles
# ------------------------------------------------------------
# Los volúmenes y montos reales de campo vienen formateados ("1.250,75 kg",
# "12,450.50 kWh"). Antes ambos formatos devolvían None y la fila entera se
# perdía como "sin columna numérica".
# ============================================================
def test_numero_con_separador_de_miles():
    from engine.parse_archivo import _a_numero

    assert _a_numero("1.234,56") == 1234.56   # ES/PE
    assert _a_numero("1,234.56") == 1234.56   # EN
    assert _a_numero("12 345,6") == 12345.6
    assert _a_numero("1,5") == 1.5
    assert _a_numero("57.4") == 57.4
    assert _a_numero("-12,5") == -12.5
    assert _a_numero("abc") is None
    assert _a_numero("") is None


# ============================================================
# PDF / DOCX — el contexto es LA LÍNEA, nunca una ventana de caracteres
# ------------------------------------------------------------
# Regresión de un bug real de atribución: con una ventana fija de 60
# caracteres, un "85,50 litros" de gasohol se clasificaba como electricidad
# porque el renglón siguiente decía "energía", y unos galones de diésel
# tomaban el factor por litro. Atribuir mal es peor que no contar.
# ============================================================
def _pdf_de_prueba(destino):
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    c = canvas.Canvas(str(destino), pagesize=A4)
    y = 800
    for linea in [
        "GRIFO EL SOL S.A.C. RUC 20512345678",
        "DIESEL B5 S-50 automotriz 1,250.00 galones S/ 18,750.00",
        "GASOHOL 90 plus 85,50 litros S/ 1,120.00",
        "Consumo de energia del local: 2.340,25 kWh",
    ]:
        c.drawString(50, y, linea)
        y -= 20
    c.save()


def test_pdf_atribuye_por_linea_y_respeta_la_unidad(tmp_path):
    from engine.ghg_classify import ghg_classify

    ruta = tmp_path / "factura.pdf"
    _pdf_de_prueba(ruta)
    clasificadas = ghg_classify(parsear_archivo(str(ruta), "pdf").lineas)
    por_unidad = {l.unidad: l for l in clasificadas}

    # Galones de diésel -> factor por GALÓN, no el de litros.
    assert por_unidad["gal"].factor_asignado == "dieselGalon"
    # El gasohol tiene su propio factor (distinto del diésel) y no debe
    # heredar "electricidad" del renglón de abajo.
    assert por_unidad["L"].factor_asignado == "gasohol84"
    assert por_unidad["kWh"].factor_asignado == "electricidadSEIN"


def test_docx_lee_parrafos_y_tablas(tmp_path):
    from docx import Document
    from engine.ghg_classify import ghg_classify

    doc = Document()
    doc.add_paragraph("Consumo de electricidad de la camara: 12,450.50 kWh")
    doc.add_paragraph("Se cargo combustible diesel B5: 340 litros")
    tabla = doc.add_table(rows=1, cols=2)
    tabla.cell(0, 0).text = "Urea aplicada"
    tabla.cell(0, 1).text = "1.250,75 kg"
    ruta = tmp_path / "orden.docx"
    doc.save(str(ruta))

    clasificadas = ghg_classify(parsear_archivo(str(ruta), "docx").lineas)
    factores = {l.factor_asignado for l in clasificadas}
    assert factores == {"electricidadSEIN", "dieselLitro", "ureaProduccion"}
    assert all(l.estado == "leido" for l in clasificadas)


def test_pdf_escaneado_sin_texto_da_motivo_claro(tmp_path):
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from engine.parse_archivo import ErrorArchivo
    import pytest

    ruta = tmp_path / "escaneado.pdf"
    c = canvas.Canvas(str(ruta), pagesize=A4)
    c.showPage()  # una página REAL pero sin capa de texto, como un escaneo
    c.save()

    with pytest.raises(ErrorArchivo) as exc:
        parsear_archivo(str(ruta), "pdf")
    assert "escaneado" in str(exc.value).lower()
