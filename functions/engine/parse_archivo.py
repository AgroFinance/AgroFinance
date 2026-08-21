"""AgroFinance — Puerto a Python de src/lib/parsing/parseArchivo.ts

Convierte un archivo (.xlsx/.csv ya soportados; .xml UBL 2.1 soportado
por separado en parsear_ubl) en LineaLeida[], el mismo formato que
consume ghg_classify. Corre en la Cloud Function, sobre el archivo ya
descargado de Cloud Storage — a diferencia del cliente, aquí no hay
límite de tiempo de un hilo de UI, pero se mantienen las mismas reglas:
una hoja/fila oculta se marca, no se descarta en silencio; un archivo
ilegible levanta un error con motivo, nunca se traga en silencio.

Alcance de este puerto: .xlsx, .csv y .xml (SUNAT UBL 2.1) — los tres
formatos que de hecho llegan hoy por Configuración/Analizar Datos y
por /upload (que es sobre todo un lector de facturas XML). .xls
(binario legado) y .ods quedan fuera por ahora — openpyxl no los lee —
y levantan ErrorArchivo con un mensaje explícito en vez de fingir
soporte.
"""

from __future__ import annotations

import csv
import math
import re
import xml.etree.ElementTree as ET
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from .ghg_classify import LineaLeida


class ErrorArchivo(Exception):
    pass


# ------------------------------------------------------------
# Deteccion de unidad por sufijo del nombre de columna — mismos
# patrones que SUFIJOS_UNIDAD en parseArchivo.ts, en el mismo orden.
# ------------------------------------------------------------
_SUFIJOS_UNIDAD: list[tuple[re.Pattern, str]] = [
    (re.compile(r"_?(kwh)$", re.I), "kWh"),
    (re.compile(r"_?(mwh)$", re.I), "MWh"),
    (re.compile(r"_?(gal|galones)$", re.I), "gal"),
    (re.compile(r"_?(m3|m³|metros?_?cubicos?)$", re.I), "m3"),
    (re.compile(r"_?(lt|ltr|litros?|l)$", re.I), "L"),
    (re.compile(r"_?(kg|kilos)$", re.I), "kg"),
    (re.compile(r"_?(tn|ton|toneladas?|t)$", re.I), "t"),
    (re.compile(r"_?(km)$", re.I), "km"),
    (re.compile(r"_?(u|und|unidades?)$", re.I), "u"),
    (re.compile(r"_?(pct|porcentaje|%)$", re.I), "%"),
]


def unidad_de_columna(col: str) -> str:
    for patron, unidad in _SUFIJOS_UNIDAD:
        if patron.search(col):
            return unidad
    return ""


def _a_numero(v) -> Optional[float]:
    if isinstance(v, (int, float)) and not isinstance(v, bool):
        return float(v) if math.isfinite(v) else None
    if not isinstance(v, str):
        return None
    limpio = v.replace(" ", "").replace(" ", "")
    if limpio == "":
        return None
    # Separador de miles: si aparecen punto Y coma, el decimal es el que está
    # más a la derecha ("1.234,56" ES/PE y "1,234.56" EN) y el otro se
    # descarta. Antes ambos formatos devolvían None y la fila se perdía como
    # "sin columna numérica", que es justo lo que pasa con los volúmenes y
    # montos reales de campo.
    if "." in limpio and "," in limpio:
        if limpio.rfind(",") > limpio.rfind("."):
            limpio = limpio.replace(".", "").replace(",", ".")
        else:
            limpio = limpio.replace(",", "")
    elif "," in limpio:
        limpio = limpio.replace(",", ".")
    try:
        return float(limpio)
    except ValueError:
        return None


@dataclass
class ResultadoParseo:
    lineas: list[LineaLeida]
    hojas: list[str]
    columnas: list[str]
    filas_preview: list[list]


# ------------------------------------------------------------
# XLSX vía openpyxl
# ------------------------------------------------------------
def _parsear_xlsx(ruta: Path) -> ResultadoParseo:
    from openpyxl import load_workbook

    wb = load_workbook(ruta, data_only=True)
    if not wb.sheetnames:
        raise ErrorArchivo("El libro no contiene hojas legibles")

    lineas: list[LineaLeida] = []
    hojas: list[str] = []
    columnas: list[str] = []
    filas_preview: list[list] = []

    for nombre_hoja in wb.sheetnames:
        ws = wb[nombre_hoja]
        hoja_oculta = ws.sheet_state != "visible"
        hojas.append(f"{nombre_hoja} (oculta)" if hoja_oculta else nombre_hoja)

        filas_iter = ws.iter_rows(values_only=False)
        try:
            encabezado_row = next(filas_iter)
        except StopIteration:
            continue
        cols = [str(c.value) if c.value is not None else "" for c in encabezado_row]
        if not any(cols):
            continue

        primera_hoja_con_datos = not columnas
        if primera_hoja_con_datos:
            columnas = cols

        for i, fila in enumerate(filas_iter, start=2):
            fila_oculta = hoja_oculta or bool(ws.row_dimensions[i].hidden)
            valores_fila = [c.value for c in fila]

            if primera_hoja_con_datos and len(filas_preview) < 12:
                filas_preview.append([v if v is not None else "" for v in valores_fila])

            for col_nombre, valor_crudo in zip(cols, valores_fila):
                if not col_nombre:
                    continue
                valor = _a_numero(valor_crudo)
                if valor is None:
                    continue
                lineas.append(LineaLeida(
                    id=f"{nombre_hoja}!{i}:{col_nombre}",
                    campo_leido=col_nombre,
                    valor=valor,
                    unidad=unidad_de_columna(col_nombre),
                    hoja=nombre_hoja,
                    fila=i,
                    oculto=fila_oculta or None,
                ))

    if not lineas:
        raise ErrorArchivo("No se encontró ninguna columna numérica que pueda leerse como consumo")

    return ResultadoParseo(lineas, hojas, columnas, filas_preview)


# ------------------------------------------------------------
# CSV vía csv estándar (misma lógica de columnas que el XLSX,
# sin concepto de hoja/fila oculta)
# ------------------------------------------------------------
def _parsear_csv(ruta: Path) -> ResultadoParseo:
    with open(ruta, newline="", encoding="utf-8-sig") as fh:
        # Excel en configuración regional ES/PE exporta CSV con ';' (la
        # coma es el separador decimal ahí) — sniffea el delimitador real
        # en vez de asumir ',', o un archivo válido cae en el mismo error
        # que uno sin columnas numéricas.
        muestra = fh.read(4096)
        fh.seek(0)
        try:
            dialecto = csv.Sniffer().sniff(muestra, delimiters=",;\t|")
        except csv.Error:
            dialecto = csv.excel
        lector = csv.reader(fh, dialecto)
        try:
            cols = next(lector)
        except StopIteration:
            raise ErrorArchivo("El archivo no contiene hojas legibles")

        lineas: list[LineaLeida] = []
        filas_preview: list[list] = []

        for i, fila in enumerate(lector, start=2):
            if len(filas_preview) < 12:
                filas_preview.append(list(fila))
            for col_nombre, valor_crudo in zip(cols, fila):
                if not col_nombre:
                    continue
                valor = _a_numero(valor_crudo)
                if valor is None:
                    continue
                lineas.append(LineaLeida(
                    id=f"CSV!{i}:{col_nombre}",
                    campo_leido=col_nombre,
                    valor=valor,
                    unidad=unidad_de_columna(col_nombre),
                    hoja="CSV",
                    fila=i,
                ))

    if not lineas:
        raise ErrorArchivo("No se encontró ninguna columna numérica que pueda leerse como consumo")

    return ResultadoParseo(lineas, ["CSV"], cols, filas_preview)


# ------------------------------------------------------------
# PDF y DOCX — documentos que NO vienen en planilla
# ------------------------------------------------------------
# Las facturas de grifo y los recibos de luz llegan casi siempre en PDF, y
# las actas/órdenes de trabajo en DOCX. No traen una grilla de columnas, así
# que no se puede usar la lógica de "columna con sufijo de unidad": acá la
# señal es el par CANTIDAD + UNIDAD dentro del texto corrido
# ("3,400.00 GALONES DE DIESEL B5", "Consumo del mes: 12 450 kWh").
#
# Se extrae ese par y se usa el texto de alrededor como campo_leido, que es
# lo que ghg_classify necesita para reconocer el factor. Igual que en el
# resto del motor: lo que no se reconoce NO se inventa.

# Unidades que el catálogo de factores sabe convertir en emisión. El orden
# importa: las más largas primero, para que "galones" gane sobre "gal".
_UNIDADES_TEXTO: list[tuple[str, str]] = [
    ("kilowatt-hora", "kWh"), ("kilovatio-hora", "kWh"), ("kilowatt hora", "kWh"),
    ("megawatt-hora", "MWh"), ("megavatio-hora", "MWh"),
    ("galones", "gal"), ("galon", "gal"), ("galón", "gal"),
    ("litros", "L"), ("litro", "L"),
    ("kilogramos", "kg"), ("kilogramo", "kg"), ("kilos", "kg"), ("kilo", "kg"),
    ("toneladas", "t"), ("tonelada", "t"),
    ("unidades", "u"), ("unidad", "u"),
    ("kwh", "kWh"), ("mwh", "MWh"), ("gal", "gal"), ("ltr", "L"),
    ("kg", "kg"), ("km", "km"), ("lt", "L"),
]

# Número con separadores de miles/decimales, seguido (opcionalmente con
# palabras en medio) de una unidad conocida.
_NUM = r"\d{1,3}(?:[.,\s]\d{3})*(?:[.,]\d+)?|\d+(?:[.,]\d+)?"
_RE_CANTIDAD_UNIDAD = re.compile(
    rf"({_NUM})\s*({'|'.join(re.escape(u) for u, _ in _UNIDADES_TEXTO)})\b",
    re.I,
)

# Layout de factura peruana: la unidad viaja en la DESCRIPCIÓN y la cantidad
# vive en su propia columna, a la derecha —
#   "Diesel B5 - Galones            180    15.80   2844.00"
#   "Consumo energia activa (kWh) - Suministro 1054877  21430  0.68  14572.40"
# El patrón número→unidad no ve nada acá porque el orden está invertido, y
# por eso las dos facturas que sostienen la promesa del producto (diésel y
# electricidad) se rechazaban enteras con "no se encontró ninguna cantidad".
_RE_UNIDAD_LUEGO_NUMEROS = re.compile(
    rf"({'|'.join(re.escape(u) for u, _ in _UNIDADES_TEXTO)})\b(?![^\W\d_])(.*)$",
    re.I,
)
def _cantidad_de_renglon(cola: str) -> tuple[float, str] | None:
    """Elige cuál de los números a la derecha de la unidad es la CANTIDAD.

    No se puede tomar el primero: en el recibo de luz el primer número es el
    N° de suministro (1054877), no el consumo (21430). Lo que sí distingue a
    la cantidad es la aritmética del renglón — cantidad × precio = importe—,
    así que se busca ese trío y se devuelve el primer factor. Sin trío que
    cuadre no se adivina: se devuelve None y el renglón queda fuera, que es
    la regla del motor (lo que no se reconoce no se inventa).

    Se tokeniza por espacios en vez de buscar _NUM suelto porque ese patrón
    admite el espacio como separador de miles, y sobre un renglón de factura
    ya normalizado partía "2844.00" en "284" y "4.00" — con los importes
    rotos la multiplicación nunca cuadraba. Cada columna es un token.
    """
    numeros = [v for v in (_a_numero(t) for t in cola.split()) if v is not None]
    for i in range(len(numeros) - 2):
        cantidad, precio, importe = numeros[i], numeros[i + 1], numeros[i + 2]
        if cantidad <= 0 or precio <= 0 or importe <= 0:
            continue
        # Tolerancia de 1% + 0.01: los importes vienen redondeados a 2 decimales.
        if abs(cantidad * precio - importe) <= max(0.01, importe * 0.01):
            return cantidad, f"{cantidad} x {precio} = {importe}"
    return None


def _lineas_desde_texto(texto: str, hoja: str) -> list[LineaLeida]:
    """Busca pares cantidad+unidad usando LA LÍNEA como contexto.

    El contexto es la línea completa, nunca una ventana de N caracteres: en
    una factura el concepto y su cantidad van en el mismo renglón, y una
    ventana fija se mete en el renglón vecino y le adjudica el consumo al
    ítem equivocado (un "85,50 litros" de gasohol terminaba clasificado como
    electricidad porque 60 caracteres más allá decía "energía"). Además, al
    conservar la línea entera el texto sigue conteniendo la palabra de la
    unidad ("galones"), que es lo que necesita la regla dieselGalon para
    ganarle a la de litros.

    Se prueban dos órdenes, en este orden de preferencia:
      1. número → unidad  ("3 400 litros de diésel"), el caso directo.
      2. unidad → números ("Diesel B5 - Galones  180  15.80  2844.00"), el
         layout de factura, resuelto por la aritmética del renglón.
    El (2) solo se evalúa si el (1) no encontró nada EN ESE renglón, para no
    duplicar el mismo consumo con dos lecturas distintas.
    """
    lineas: list[LineaLeida] = []
    n = 0
    for fila, linea in enumerate(texto.splitlines(), start=1):
        limpia = " ".join(linea.split())
        if not limpia:
            continue

        encontrados: list[tuple[float, str, str]] = [
            (valor, next(u for txt, u in _UNIDADES_TEXTO if txt.lower() == m.group(2).lower()), m.group(0))
            for m in _RE_CANTIDAD_UNIDAD.finditer(limpia)
            if (valor := _a_numero(m.group(1))) is not None
        ]

        if not encontrados:
            m = _RE_UNIDAD_LUEGO_NUMEROS.search(limpia)
            if m and (hallazgo := _cantidad_de_renglon(m.group(2))):
                valor, crudo = hallazgo
                unidad = next(u for txt, u in _UNIDADES_TEXTO if txt.lower() == m.group(1).lower())
                encontrados = [(valor, unidad, crudo)]

        for valor, unidad, crudo in encontrados:
            n += 1
            lineas.append(LineaLeida(
                id=f"{hoja}!linea-{fila}-{n}",
                campo_leido=limpia,
                valor=valor,
                unidad=unidad,
                hoja=hoja,
                fila=fila,
                crudo=crudo,
            ))
    return lineas


def _parsear_pdf(ruta: Path) -> ResultadoParseo:
    import pdfplumber

    partes: list[str] = []
    filas_preview: list[list] = []
    try:
        with pdfplumber.open(ruta) as pdf:
            if not pdf.pages:
                raise ErrorArchivo("El PDF no tiene páginas legibles")
            for pagina in pdf.pages:
                partes.append(pagina.extract_text() or "")
                for tabla in pagina.extract_tables() or []:
                    for fila in tabla:
                        celdas = [(c or "").strip() for c in fila]
                        if not any(celdas):
                            continue
                        # La fila de una tabla se aplana a texto: así el mismo
                        # buscador de cantidad+unidad sirve para tablas y para
                        # texto corrido, sin asumir que la tabla tiene encabezado.
                        partes.append(" ".join(celdas))
                        if len(filas_preview) < 12:
                            filas_preview.append(celdas)
    except ErrorArchivo:
        raise
    except Exception as exc:  # noqa: BLE001 — se declara el motivo, no se traga
        raise ErrorArchivo(f"No se pudo leer el PDF: {exc}")

    texto = "\n".join(partes)
    if not texto.strip():
        raise ErrorArchivo(
            "El PDF no contiene texto extraíble — parece escaneado como imagen. "
            "Vuelve a exportarlo desde el sistema que lo emitió, o súbelo en XML/Excel."
        )

    # Un PDF legible SIN consumos no es un archivo inválido: un certificado
    # GlobalGAP o una constancia no tienen litros ni kWh y nunca los tendrán.
    # Antes se rechazaban con ErrorArchivo y la carpeta entera se llenaba de
    # errores rojos por documentos que en realidad se leyeron perfecto — se
    # devuelven con cero líneas, que es la verdad: leído, sin consumos.
    lineas = _lineas_desde_texto(texto, "PDF")

    if not filas_preview:
        filas_preview = [[l.campo_leido, l.valor, l.unidad] for l in lineas[:12]]
    return ResultadoParseo(lineas, ["PDF"], ["Detalle detectado", "Cantidad", "Unidad"], filas_preview)


def _parsear_docx(ruta: Path) -> ResultadoParseo:
    from docx import Document

    try:
        doc = Document(str(ruta))
    except Exception as exc:  # noqa: BLE001
        raise ErrorArchivo(f"No se pudo leer el documento Word: {exc}")

    partes = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    filas_preview: list[list] = []
    for tabla in doc.tables:
        for fila in tabla.rows:
            celdas = [c.text.strip() for c in fila.cells]
            if not any(celdas):
                continue
            partes.append(" ".join(celdas))
            if len(filas_preview) < 12:
                filas_preview.append(celdas)

    texto = "\n".join(partes)
    if not texto.strip():
        raise ErrorArchivo("El documento Word está vacío o no tiene texto legible")

    # Mismo criterio que el PDF: una minuta de auditoría sin consumos se leyó
    # bien, solo que no aporta emisiones. Cero líneas, no error.
    lineas = _lineas_desde_texto(texto, "DOCX")

    if not filas_preview:
        filas_preview = [[l.campo_leido, l.valor, l.unidad] for l in lineas[:12]]
    return ResultadoParseo(lineas, ["DOCX"], ["Detalle detectado", "Cantidad", "Unidad"], filas_preview)


# ------------------------------------------------------------
# XML SUNAT UBL 2.1 — mismos namespaces y misma lógica que parsearUBL
# en parseArchivo.ts (getElementsByTagName ahí == .iter() acá: ambos
# buscan en TODO el subárbol, no solo hijos directos).
# ------------------------------------------------------------
_NS = {
    "cac": "urn:oasis:names:specification:ubl:schema:xsd:CommonAggregateComponents-2",
    "cbc": "urn:oasis:names:specification:ubl:schema:xsd:CommonBasicComponents-2",
}

_UNIDAD_UBL = {"LTR": "L", "GLL": "gal", "KWH": "kWh", "MWH": "MWh", "KGM": "kg", "TNE": "t", "NIU": "u", "ZZ": ""}


def _tag(prefix: str, local: str) -> str:
    return f"{{{_NS[prefix]}}}{local}"


def _primero(elem: ET.Element, prefix: str, local: str) -> Optional[ET.Element]:
    return next(elem.iter(_tag(prefix, local)), None)


def _texto(elem: ET.Element, prefix: str, local: str) -> str:
    nodo = _primero(elem, prefix, local)
    return (nodo.text or "").strip() if nodo is not None else ""


@dataclass
class CabeceraUBL:
    invoice_id: str
    ruc: str
    proveedor: str
    fecha: str
    moneda: str
    monto_total: Optional[float]


@dataclass
class ResultadoUBL(ResultadoParseo):
    cabecera: CabeceraUBL = None  # type: ignore[assignment]


def _parsear_ubl(ruta: Path, nombre: str) -> ResultadoUBL:
    try:
        root = ET.parse(ruta).getroot()
    except ET.ParseError as exc:
        raise ErrorArchivo(f"El XML no está bien formado y no pudo interpretarse: {exc}")

    lineas_xml = list(root.iter(_tag("cac", "InvoiceLine")))
    raiz_ids = list(root.iter(_tag("cbc", "ID")))

    pago = _primero(root, "cbc", "PayableAmount")
    cabecera = CabeceraUBL(
        invoice_id=(raiz_ids[0].text.strip() if raiz_ids and raiz_ids[0].text else nombre),
        ruc=(raiz_ids[1].text.strip() if len(raiz_ids) > 1 and raiz_ids[1].text else "—"),
        proveedor=_texto(root, "cbc", "RegistrationName") or "—",
        fecha=_texto(root, "cbc", "IssueDate") or "—",
        moneda=(pago.get("currencyID") if pago is not None and pago.get("currencyID") else "PEN"),
        monto_total=_a_numero(_texto(root, "cbc", "PayableAmount")),
    )

    lineas: list[LineaLeida] = []
    filas_preview: list[list] = []
    for i, nodo in enumerate(lineas_xml):
        qty_nodo = _primero(nodo, "cbc", "InvoicedQuantity")
        descripcion = _texto(nodo, "cbc", "Description") or f"Ítem {i + 1}"
        unit_code = qty_nodo.get("unitCode", "") if qty_nodo is not None else ""
        valor = _a_numero(qty_nodo.text) if qty_nodo is not None else None
        unidad = _UNIDAD_UBL.get(unit_code, unit_code)
        lineas.append(LineaLeida(
            id=f"UBL!linea-{i + 1}", campo_leido=descripcion, valor=valor, unidad=unidad,
            hoja="Comprobante UBL 2.1", fila=i + 1,
            crudo=(qty_nodo.text.strip() if qty_nodo is not None and qty_nodo.text else None),
        ))
        if len(filas_preview) < 12:
            filas_preview.append([descripcion, valor if valor is not None else "", unidad])

    if not lineas:
        raise ErrorArchivo("El comprobante no declara líneas de detalle (cac:InvoiceLine)")

    return ResultadoUBL(
        lineas=lineas, hojas=["Comprobante UBL 2.1"],
        columnas=["Descripción del ítem", "Cantidad", "Unidad"], filas_preview=filas_preview,
        cabecera=cabecera,
    )


# ------------------------------------------------------------
# Punto de entrada
# ------------------------------------------------------------
def parsear_archivo(ruta_local: str, tipo: str) -> ResultadoParseo:
    ruta = Path(ruta_local)
    tipo = tipo.lower().lstrip(".")

    if tipo == "xlsx":
        return _parsear_xlsx(ruta)
    if tipo == "csv":
        return _parsear_csv(ruta)
    if tipo == "xml":
        return _parsear_ubl(ruta, ruta.name)
    if tipo == "pdf":
        return _parsear_pdf(ruta)
    if tipo == "docx":
        return _parsear_docx(ruta)
    if tipo == "doc":
        raise ErrorArchivo(
            "El formato .doc (Word 97-2003) no se puede leer. Ábrelo en Word y guárdalo como .docx."
        )
    if tipo in ("xls", "ods"):
        raise ErrorArchivo(
            f"Formato .{tipo} aún no soportado en el procesamiento en la nube — "
            "usa .xlsx, .csv o .xml, o súbelo desde Configuración (que sí lo lee en el navegador)."
        )
    raise ErrorArchivo(f"Formato .{tipo} no soportado")
